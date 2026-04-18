import re
import os
from docx import Document
import openpyxl
from openpyxl.styles import Alignment
from copy import copy

# 文件常見名稱 → 委辦補助官方名稱
EXPENSE_NAME_MAP = {
    '講座鐘點費':       '講座鐘點費',
    '校內講座鐘點':     '講座鐘點費',
    '鐘點費':           '講座鐘點費',
    '主持費':           '主持費、引言費',
    '引言費':           '主持費、引言費',
    '膳費':             '膳宿費',
    '膳宿費':           '膳宿費',
    '交通費':           '國內旅費、短程車資、運費',
    '旅費':             '國內旅費、短程車資、運費',
    '車資':             '國內旅費、短程車資、運費',
    '租車費':           '租車費',
    '遊覽車':           '租車費',
    '印刷費':           '印刷費',
    '材料費':           '材料費',
    '物品費':           '物品費',
    '雜支':             '雜支',
    '保險費':           '保險費',
    '場地費':           '場地使用費',
    '場地使用費':       '場地使用費',
    '設備使用費':       '設備使用費',
    '全民健康保險補充保費': '全民健康保險補充保費',
    '健保補充費':       '全民健康保險補充保費',
}


class BudgetProcessor:
    def __init__(self, basis_path, doc_path, template_path,
                 school_name='', fiscal_year=''):
        self.basis_path = basis_path
        self.doc_path = doc_path
        self.template_path = template_path
        self.school_name = school_name
        self.fiscal_year = fiscal_year

    # ─────────────────────────────────────────
    # Public entry point
    # ─────────────────────────────────────────
    def process(self):
        basis_content = self._read_basis()
        doc_info      = self._parse_docx()
        template_info = self._read_template_info()

        skill_content = self._generate_skill(basis_content, doc_info, template_info)
        excel_path    = self._generate_excel(doc_info, template_info)

        plan_code = doc_info.get('plan_code', '?')
        part1     = doc_info['sections']['PART1']
        part2     = doc_info['sections']['PART2']
        both      = doc_info['sections']['BOTH']

        return {
            'skill_content': skill_content,
            'excel_path':    excel_path,
            'summary': {
                'plan_code':   plan_code,
                'part1_code':  part1['code'] if part1 else None,
                'part2_code':  part2['code'] if part2 else None,
                'both_codes':  [s['code'] for s in both],
                'sheets':      template_info.get('sheets', []),
            }
        }

    # ─────────────────────────────────────────
    # Parsers
    # ─────────────────────────────────────────
    def _read_basis(self):
        with open(self.basis_path, 'r', encoding='utf-8') as f:
            return f.read()

    def _parse_docx(self):
        doc = Document(self.doc_path)
        result = {
            'plan_code': None,
            'sections': {'PART1': None, 'PART2': None, 'BOTH': []}
        }

        current_type  = None
        current_code  = None
        current_paras = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # PART1 / PART2 markers (fixed)
            m = re.match(r'^PART(\d+)[:\s]*([A-Z]-\d+-\d+)?', text)
            if m:
                self._flush(result, current_type, current_code, current_paras)
                current_type  = 'PART1' if m.group(1) == '1' else 'PART2'
                current_code  = m.group(2)
                current_paras = []
                if current_code and not result['plan_code']:
                    result['plan_code'] = current_code[0]
                continue

            # Standalone plan code, e.g. "B-2-1"
            if re.match(r'^[A-Z]-\d+-\d+$', text):
                self._flush(result, current_type, current_code, current_paras)
                current_type  = 'STANDALONE'
                current_code  = text
                current_paras = []
                if not result['plan_code']:
                    result['plan_code'] = text[0]
                continue

            current_paras.append(text)

        self._flush(result, current_type, current_code, current_paras)
        return result

    def _flush(self, result, section_type, code, paragraphs):
        if not section_type or not code:
            return
        full_text = '\n'.join(paragraphs)
        data = {
            'code':               code,
            'paragraphs':         paragraphs,
            'expenses':           self._extract_expenses(paragraphs),
            'has_every_semester':  '每學期' in full_text,
            'has_upper_semester':  '上學期' in full_text and '每學期' not in full_text,
            'has_lower_semester':  '下學期' in full_text and '每學期' not in full_text,
            'full_text':          full_text,
        }
        if section_type == 'PART1':
            result['sections']['PART1'] = data
        elif section_type == 'PART2':
            result['sections']['PART2'] = data
        else:
            result['sections']['BOTH'].append(data)

    def _extract_expenses(self, paragraphs):
        expenses = []
        for para in paragraphs:
            # Pattern: 費用名稱 金額 * 數量 (e.g., 講座鐘點費2000*3)
            for raw_name, price_str, qty_str in re.findall(
                r'([\u4e00-\u9fff]{2,10}[費鐘])\s*(\d+(?:\.\d+)?)\s*(?:元)?\s*[*×xX]\s*(\d+(?:\.\d+)?)',
                para
            ):
                expenses.append({
                    'raw_name': raw_name,
                    'name':     EXPENSE_NAME_MAP.get(raw_name, raw_name),
                    'price':    float(price_str),
                    'qty':      float(qty_str),
                    'source':   para[:80],
                })

            # Pattern: 遊覽車 N 部 金額元
            for price_str in re.findall(r'遊覽車(?:一|[\d]+)?部.*?(\d+)元', para):
                expenses.append({
                    'raw_name': '遊覽車',
                    'name':     '租車費',
                    'price':    float(price_str),
                    'qty':      1,
                    'source':   para[:80],
                })
        return expenses

    def _read_template_info(self):
        try:
            wb = openpyxl.load_workbook(self.template_path)
            return {'sheets': wb.sheetnames, 'path': self.template_path}
        except Exception:
            return {'sheets': [], 'path': self.template_path}

    # ─────────────────────────────────────────
    # Skill generator
    # ─────────────────────────────────────────
    def _generate_skill(self, basis_content, doc_info, template_info):
        plan_code  = doc_info.get('plan_code', 'X')
        part1      = doc_info['sections']['PART1']
        part2      = doc_info['sections']['PART2']
        both       = doc_info['sections']['BOTH']

        part1_code     = part1['code'] if part1 else '未偵測到'
        part2_code     = part2['code'] if part2 else '未偵測到'
        both_codes_str = '、'.join(s['code'] for s in both) or '無'

        # Aggregate detected expenses
        seen = {}
        for sec in ([part1] if part1 else []) + ([part2] if part2 else []) + both:
            for exp in sec.get('expenses', []):
                n = exp['name']
                seen[n] = seen.get(n, 0) + 1
        expense_lines = '\n'.join(
            f'- {n}（文件中出現 {c} 次）' for n, c in seen.items()
        ) or '- 請手動解析文件'

        target_sheets = [s for s in template_info['sheets'] if plan_code in s]
        sheets_lines = '\n'.join(f'- {s}' for s in target_sheets) or '- 請依代碼手動確認'

        school_line = f'- 學校：{self.school_name}' if self.school_name else ''
        year_line   = f'- 會計年度：{self.fiscal_year}' if self.fiscal_year else '- 會計年度：（未提供，請補填）'

        return f'''\
---
name: budget-compiler
description: 當使用者要將計畫經費填入經費編列表時使用此 skill。觸發詞包括「填寫經費」、「編列經費」、「填入經費編列表」、「幫我寫經費」、「整理預算」。此 skill 參考委辦補助.md（國教署補助標準）與範例.docx（計畫活動說明），將資料整理後寫入經費編列表.xlsx。
argument-hint: [委辦補助.md路徑] [範例.docx路徑] [經費編列表.xlsx路徑]
allowed-tools: [Read, Bash, Glob]
---

# 經費編列表填寫工具

## 本計畫基本資料（由上傳檔案自動偵測）

{school_line}
{year_line}
- 子計畫代碼：**{plan_code}**
- PART1 代碼（8-12月）：{part1_code}
- PART2 代碼（1-7月）：{part2_code}
- 每學期活動代碼：{both_codes_str}

## 偵測到的費用類型

{expense_lines}

## 目標工作表

{sheets_lines}

---

## 工作目標

依據 **委辦補助.md**（國教署標準）與 **範例.docx**（計畫活動說明），
將經費正確填入 **經費編列表.xlsx** 對應的子計畫工作表。

## 檔案路徑

若使用者有提供路徑，以使用者提供的為準；否則預設從當前工作目錄讀取：
- 委辦補助.md
- 範例.docx
- 經費編列表.xlsx

---

## 執行步驟

### STEP 1：讀取補助標準

使用 Read 工具讀取 `委辦補助.md`，**完整閱讀全文**，建立以下對照表：

| 擷取項目 | 說明 |
|----------|------|
| **官方名稱** | 各費用項目的正式名稱（不可自創或縮寫） |
| **單價範圍** | 每節/每人次/每份等的金額上下限 |
| **單位** | 計算基礎（節、人次、式、份、臺…） |

> 委辦補助.md 為唯一標準，不得依常識自行假設。

---

### STEP 2：讀取範例.docx

```bash
python3 -c "
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
doc = Document('範例.docx')
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f'[{{i}}] {{para.text}}')
for t_idx, table in enumerate(doc.tables):
    print(f'-- Table {{t_idx}} --')
    for row in table.rows:
        print([c.text.strip() for c in row.cells])
"
```

#### 2-1 文件結構（固定標記）

**學期定義**：上學期 = 8-12月、下學期 = 1-7月、每學期 = 上學期及下學期皆編列

| 標記/關鍵詞 | 對應學期 | 目標工作表 |
|-------------|----------|------------|
| `PART1` 或描述含「**上學期**」 | 上學期（8-12月） | `{plan_code}(8-12月)` |
| `PART2` 或描述含「**下學期**」 | 下學期（1-7月） | `{plan_code}(1-7月)` |
| 描述含「**每學期**」 | 上學期＋下學期 | `{plan_code}(8-12月)` **及** `{plan_code}(1-7月)` |

**子計畫整合規則**：相同首字母的所有子計畫（如 {plan_code}-1-1、{plan_code}-1-2、{plan_code}-2-1）**整合至同一張工作表**，在「說明用途」欄以代碼分段標註。

#### 2-2 學期編列規則

| 文件描述 | 處理方式 |
|----------|----------|
| 段落標記 `PART1` 或含「上學期」 | 只填 **`{plan_code}(8-12月)`** 工作表 |
| 段落標記 `PART2` 或含「下學期」 | 只填 **`{plan_code}(1-7月)`** 工作表 |
| 描述中含「**每學期**」 | **同時填入兩張工作表**，各自以當學期數量計算 |
| 無標記也無學期關鍵詞 | 詢問使用者確認 |

「每學期 N 場次」→ 每張工作表各填 N 場次（不跨學期加總）

---

### STEP 3：確認工作表

```bash
python3 -c "
import sys; sys.stdout.reconfigure(encoding='utf-8')
import openpyxl
wb = openpyxl.load_workbook('經費編列表.xlsx')
print(wb.sheetnames)
"
```

---

### STEP 4：整理填寫計畫

#### 欄位規則

| 欄位 | 規則 |
|------|------|
| **名稱**（B） | 完全符合委辦補助.md官方名稱，不可自創 |
| **單位**（C） | 文件有載明依文件；未載明**留空** |
| **數量**（D） | 各子計畫加總；多來源用公式 `=數字+數字+...` |
| **單價**（E） | 依委辦補助.md，**單位仟元**（2000元→`2`，160元→`0.16`） |
| **總價**（F） | 公式 `=D行*E行` |
| **說明用途**（G） | 各子計畫分行；**不同代碼間強制換行**，設 wrap_text=True |
| **經費來源**（H） | 通常留空 |

#### 說明用途格式（強制換行）

```
B-X-1 社群名稱：
　活動說明1
　活動說明2
B-X-2 社群名稱：
　活動說明1
```

#### 資本門判斷

| 條件 | 歸屬 |
|------|------|
| 單件單價 ≤ 10,000元 | 經常門 → 業務費 |
| 單件單價 > 10,000元 | 資本門 → 設備及投資 |

#### 排列順序（經常門）

講座鐘點費 → 膳宿費 → 雜支 → 國內旅費、短程車資、運費 → 全民健康保險補充保費 → 租車費 → 物品費 → 印刷費 → 材料費

---

### STEP 5：寫入 Excel

```bash
python3 << 'PYEOF'
import sys
sys.stdout.reconfigure(encoding='utf-8')
import openpyxl
from openpyxl.styles import Alignment

wb = openpyxl.load_workbook('經費編列表.xlsx')

def unmerge_clear(ws, start, end):
    for mr in list(ws.merged_cells.ranges):
        if mr.min_row <= end and mr.max_row >= start:
            ws.unmerge_cells(str(mr))
    for r in range(start, end+1):
        for c in range(1, 9):
            ws.cell(row=r, column=c).value = None

def write_sheet(ws, items_regular, items_capital):
    START = 8
    unmerge_clear(ws, START, 35)
    for i, (name, unit, qty, price, desc) in enumerate(items_regular):
        r = START + i
        ws.cell(row=r, column=1).value = '業\\n務\\n費' if i == 0 else None
        ws.cell(row=r, column=2).value = name
        ws.cell(row=r, column=3).value = unit or None
        ws.cell(row=r, column=4).value = qty
        ws.cell(row=r, column=5).value = price
        ws.cell(row=r, column=6).value = f'=D{{r}}*E{{r}}'
        c = ws.cell(row=r, column=7, value=desc)
        c.alignment = Alignment(wrap_text=True, vertical='top')
        ws.cell(row=r, column=8).value = ''
    sub = START + len(items_regular)
    ws.cell(row=sub, column=2).value = '小計'
    ws.cell(row=sub, column=6).value = f'=SUM(F{{START}}:F{{sub-1}})'
    cap_h = sub + 1
    ws.cell(row=cap_h, column=1).value = '(二)資本門'
    cap_s = cap_h + 1
    if items_capital:
        for i, (name, unit, qty, price, desc) in enumerate(items_capital):
            r = cap_s + i
            ws.cell(row=r, column=1).value = '設\\n備\\n及\\n投\\n資' if i == 0 else None
            ws.cell(row=r, column=2).value = name
            ws.cell(row=r, column=3).value = unit or None
            ws.cell(row=r, column=4).value = qty
            ws.cell(row=r, column=5).value = price
            ws.cell(row=r, column=6).value = f'=D{{r}}*E{{r}}'
            c = ws.cell(row=r, column=7, value=desc)
            c.alignment = Alignment(wrap_text=True, vertical='top')
        cap_last = cap_s + len(items_capital) - 1
    else:
        ws.cell(row=cap_s, column=1).value = '設\\n備\\n及\\n投\\n資'
        cap_last = cap_s
    cap_sub = cap_last + 1
    ws.cell(row=cap_sub, column=1).value = '資本門小計'
    ws.cell(row=cap_sub, column=6).value = f'=SUM(F{{cap_s}}:F{{cap_last}})' if items_capital else 0
    total = cap_sub + 1
    ws.cell(row=total, column=1).value = '年度合計'
    ws.cell(row=total, column=6).value = f'=F{{sub}}+F{{cap_sub}}'

# 依 STEP 4 整理的資料填入
items_y11 = []  # (名稱, 單位, 數量or公式, 單價仟元, 說明用途)
items_y12 = []
items_cap_y11 = []
items_cap_y12 = []

sheet_names = wb.sheetnames
# 優先找新命名格式 {plan_code}(8-12月)，找不到再找舊格式
y11 = next((s for s in sheet_names if s == '{plan_code}(8-12月)'), None) \
   or next((s for s in sheet_names if '{plan_code}' in s and ('8-12' in s or 'y-1-1' in s)), None)
y12 = next((s for s in sheet_names if s == '{plan_code}(1-7月)'), None) \
   or next((s for s in sheet_names if '{plan_code}' in s and ('1-7' in s or 'y-1-2' in s)), None)

if y11: write_sheet(wb[y11], items_y11, items_cap_y11)
if y12: write_sheet(wb[y12], items_y12, items_cap_y12)

wb.save('output.xlsx')
print('完成')
PYEOF
```

### STEP 6：驗證

```bash
python3 -c "
import sys; sys.stdout.reconfigure(encoding='utf-8')
import openpyxl
wb = openpyxl.load_workbook('output.xlsx')
for s in wb.sheetnames:
    if '{plan_code}' in s:
        ws = wb[s]
        print(f'=== {{s}} ===')
        for i, row in enumerate(ws.iter_rows(values_only=True), 1):
            if any(c for c in row):
                print(f'  row {{i}}:', [str(c)[:35] if c else '' for c in row[:7]])
"
```

---

## 注意事項

1. **名稱嚴格對應**：「交通費」需改為「國內旅費、短程車資、運費」
2. **單位仟元**：2,000元→`2`；160元→`0.16`；12,000元→`12`
3. **相同費用合併**：跨子計畫同類費用合一行
4. **資本門界線**：單件單價 > 10,000元才放資本門
5. **全民健保**：`=ROUND(F講座鐘點費行*0.0211,3)`
6. **每學期**：含「每學期」的段落同時填入兩張工作表，各自計算
7. **說明用途換行**：不同代碼間必須 `\\n` 換行，設 `wrap_text=True`

---

## 完成後回報

- 已填入的工作表名稱
- 各費用項目摘要（名稱、數量、單價、總價）
- 進入資本門的項目
- 留空欄位清單（需人工補填）
'''

    # ─────────────────────────────────────────
    # Excel generator
    # ─────────────────────────────────────────
    def _generate_excel(self, doc_info, template_info):
        """
        Generate a new Excel file based on the template and parsed doc.
        Creates new.xlsx alongside the template.
        """
        try:
            plan_code = doc_info.get('plan_code', 'B')
            src_path  = template_info['path']
            src_wb    = openpyxl.load_workbook(src_path)
            sheets    = src_wb.sheetnames

            # 統一使用新命名格式 {plan_code}(8-12月) / {plan_code}(1-7月)
            y11_name = f'{plan_code}(8-12月)'
            y12_name = f'{plan_code}(1-7月)'

            # Build expense items from parsed sections
            part1    = doc_info['sections']['PART1']
            part2    = doc_info['sections']['PART2']
            both_sec = doc_info['sections']['BOTH']

            items_y11 = self._build_items(part1, both_sec)
            items_y12 = self._build_items(part2, both_sec)

            # Write to a new workbook based on template
            new_wb = openpyxl.Workbook()
            new_wb.remove(new_wb.active)

            # Copy Example sheet if exists
            if 'Example' in sheets:
                self._copy_sheet(src_wb['Example'], new_wb, 'Example')

            ws = new_wb.create_sheet(title=y11_name)
            self._init_header(ws, plan_code, '8-12月', self.school_name, self.fiscal_year)
            self._write_items(ws, items_y11['regular'], items_y11['capital'])

            ws = new_wb.create_sheet(title=y12_name)
            self._init_header(ws, plan_code, '1-7月', self.school_name, self.fiscal_year)
            self._write_items(ws, items_y12['regular'], items_y12['capital'])

            out_path = os.path.join(os.path.dirname(src_path), 'generated.xlsx')
            new_wb.save(out_path)
            return out_path

        except Exception as e:
            print(f'Excel generation error: {e}')
            return None

    def _build_items(self, main_section, both_sections):
        """Aggregate expense items from a PART section + BOTH sections."""
        # Collect all expenses grouped by official name
        grouped = {}

        all_secs = ([main_section] if main_section else []) + list(both_sections)
        for sec in all_secs:
            code = sec['code']
            for exp in sec.get('expenses', []):
                name = exp['name']
                if name not in grouped:
                    grouped[name] = {
                        'name':    name,
                        'price':   exp['price'],
                        'entries': [],
                    }
                grouped[name]['entries'].append({
                    'code':   code,
                    'qty':    exp['qty'],
                    'source': exp.get('source', ''),
                })

        regular  = []
        capital  = []
        PRIORITY = ['講座鐘點費', '膳宿費', '雜支', '國內旅費、短程車資、運費',
                    '全民健康保險補充保費', '租車費', '物品費', '印刷費', '材料費']

        def sort_key(item):
            try:
                return PRIORITY.index(item['name'])
            except ValueError:
                return len(PRIORITY)

        for item in sorted(grouped.values(), key=sort_key):
            entries = item['entries']
            qty_parts = [str(int(e['qty'])) for e in entries]
            qty_formula = '=' + '+'.join(qty_parts) if len(qty_parts) > 1 else float(entries[0]['qty'])

            desc_lines = []
            for e in entries:
                desc_lines.append(f"{e['code']}：{e['source'][:60]}")
            desc = '\n'.join(desc_lines)

            price_kilo = item['price'] / 1000
            row = (item['name'], '', qty_formula, price_kilo, desc)

            if item['price'] > 10000:
                capital.append(row)
            else:
                regular.append(row)

        return {'regular': regular, 'capital': capital}

    def _init_header(self, ws, plan_code, period, school_name, fiscal_year):
        title = f'{school_name} 子計畫概算表({period})' if school_name else f'子計畫概算表({period})'
        ws.cell(row=1, column=1).value = title
        ws.cell(row=3, column=1).value = f'子計畫代碼：{plan_code}'
        ws.cell(row=4, column=7).value = '單位：仟元'
        ws.cell(row=5, column=1).value = f'{fiscal_year}會計年度概算表({period})' if fiscal_year else ''
        ws.cell(row=6, column=1).value = ''
        for col, header in enumerate(['', '名稱', '單位', '數量', '單價', '總價', '說明用途', '經費來源'], 1):
            ws.cell(row=6, column=col).value = header
        ws.cell(row=7, column=1).value = '(一)經常門'

    def _write_items(self, ws, regular, capital):
        START = 8
        for i, (name, unit, qty, price, desc) in enumerate(regular):
            r = START + i
            ws.cell(row=r, column=1).value = '業\n務\n費' if i == 0 else None
            ws.cell(row=r, column=2).value = name
            ws.cell(row=r, column=3).value = unit or None
            ws.cell(row=r, column=4).value = qty
            ws.cell(row=r, column=5).value = price
            ws.cell(row=r, column=6).value = f'=D{r}*E{r}'
            c = ws.cell(row=r, column=7, value=desc)
            c.alignment = Alignment(wrap_text=True, vertical='top')
            ws.cell(row=r, column=8).value = ''

        sub = START + len(regular)
        ws.cell(row=sub, column=2).value = '小計'
        ws.cell(row=sub, column=6).value = f'=SUM(F{START}:F{sub-1})'

        cap_h = sub + 1
        ws.cell(row=cap_h, column=1).value = '(二)資本門'
        cap_s = cap_h + 1

        if capital:
            for i, (name, unit, qty, price, desc) in enumerate(capital):
                r = cap_s + i
                ws.cell(row=r, column=1).value = '設\n備\n及\n投\n資' if i == 0 else None
                ws.cell(row=r, column=2).value = name
                ws.cell(row=r, column=3).value = unit or None
                ws.cell(row=r, column=4).value = qty
                ws.cell(row=r, column=5).value = price
                ws.cell(row=r, column=6).value = f'=D{r}*E{r}'
                c = ws.cell(row=r, column=7, value=desc)
                c.alignment = Alignment(wrap_text=True, vertical='top')
            cap_last = cap_s + len(capital) - 1
        else:
            ws.cell(row=cap_s, column=1).value = '設\n備\n及\n投\n資'
            cap_last = cap_s

        cap_sub = cap_last + 1
        ws.cell(row=cap_sub, column=1).value = '資本門小計'
        ws.cell(row=cap_sub, column=6).value = f'=SUM(F{cap_s}:F{cap_last})' if capital else 0

        total = cap_sub + 1
        ws.cell(row=total, column=1).value = '年度合計'
        ws.cell(row=total, column=6).value = f'=F{sub}+F{cap_sub}'

        sign = total + 1
        ws.cell(row=sign, column=1).value = (
            '承辦人：\u3000\u3000\u3000\u3000\u3000\u3000\u3000'
            '承辦主任：\u3000\u3000\u3000\u3000\u3000\u3000\u3000'
            '主計主任：\u3000\u3000\u3000\u3000\u3000\u3000\u3000'
            '校長：\u3000\u3000\u3000\u3000\u3000\u3000\u3000'
        )

    def _copy_sheet(self, src_ws, dst_wb, title):
        dst_ws = dst_wb.create_sheet(title=title)
        for row in src_ws.iter_rows():
            for cell in row:
                dst_cell = dst_ws.cell(row=cell.row, column=cell.column)
                dst_cell.value = cell.value
                if cell.has_style:
                    dst_cell.font       = copy(cell.font)
                    dst_cell.border     = copy(cell.border)
                    dst_cell.fill       = copy(cell.fill)
                    dst_cell.alignment  = copy(cell.alignment)
                    dst_cell.number_format = cell.number_format
        for mc in src_ws.merged_cells.ranges:
            try:
                dst_ws.merge_cells(str(mc))
            except Exception:
                pass
        for col, dim in src_ws.column_dimensions.items():
            dst_ws.column_dimensions[col].width = dim.width
        return dst_ws
